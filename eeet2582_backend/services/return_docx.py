import re
import base64
import os
from io import BytesIO
from docx.shared import Inches
from docx import Document
import boto3

from django.conf import settings
from rest_framework.response import Response
from eeet2582_backend.api.models.document_paragraph import DocumentParagraph
from eeet2582_backend.api.models.endnote import EndNote
from eeet2582_backend.api.models.heading import Heading
from eeet2582_backend.api.models.list_paragraph import ListParagraph
from eeet2582_backend.api.models.user_document import UserDocument
from eeet2582_backend.api.models.subheading import Subheading
from eeet2582_backend.api.models.document_table import DocumentTable
from eeet2582_backend.api.models.table_row import TableRow
from eeet2582_backend.api.models.row_cell import RowCell
from eeet2582_backend.api.models.document_image import DocumentImage
from eeet2582_backend.api.models.caption import Caption
from eeet2582_backend.celery import app
from eeet2582_backend.services import upload_to_s3

# Function remove invalid characters in filename
def remove_invalid(filename):
    return "".join([c for c in filename if c.isalpha() or c.isdigit() or c == ' ']).rstrip()

# Function add heading to the document
def add_headings(doc, heading_content, level):
    doc.add_heading(heading_content, level)

# Function add paragraph to the document
def add_paragraphs(doc, content):
    doc.add_paragraph(content)

# Fucntion add a list of paragraph to the document
def add_list_paragraph(doc, content):
    doc.add_paragraph(content, style='List Bullet')

# Function add table to the document
def add_table_to_doc(doc, table):
    if not table.tablerow_set.exists():
        return

    # Determine the number of columns
    number_of_columns = len(table.tablerow_set.first().rowcell_set.all())
    doc_table = doc.add_table(rows=0, cols=number_of_columns)
    doc_table.style = 'Table Grid'

    # Add rows and cells to the table
    for row in table.tablerow_set.all():
        cells = row.rowcell_set.all()
        row_cells = doc_table.add_row().cells
        for idx, cell in enumerate(cells):
            row_cells[idx].text = cell.content

# Function add image to the document
def add_image_to_doc(doc, base64_string):
    # Decode the base64 string to bytes and create an image stream
    image_data = base64.b64decode(base64_string)
    image_stream = BytesIO(image_data)

    # Add image to the document
    doc.add_picture(image_stream, width=Inches(6))

# Function add caption to the document
def add_caption_to_doc(doc, caption):
    p = doc.add_paragraph()
    p.add_run(caption).italic = True
    
class ReturnDocxService:
    @app.task
    # Function create word docx document
    def create_docx(self, user_doc_id, file_path):
        user_doc = UserDocument.objects.get(id=user_doc_id)
        # Create a new Word document
        doc = Document()
        normalized_path = os.path.normpath(file_path)
        file_name = os.path.split(normalized_path)[-1]

        # Add the title to the document
        doc_title = user_doc.document_title.title
        add_headings(doc, doc_title, level=0)

        # Fetch and sort headings and paragraphs
        headings = Heading.objects.filter(user_document=user_doc)
        list_paragraphs = ListParagraph.objects.filter(user_document=user_doc)
        paragraphs = DocumentParagraph.objects.filter(user_document=user_doc).order_by('id')
        tables = DocumentTable.objects.filter(user_document=user_doc)
        images = DocumentImage.objects.filter(user_document=user_doc)
        captions = Caption.objects.filter(user_document=user_doc)

        for para in paragraphs:
            # Filter which element has same paragraph ID
            related_headings = headings.filter(document_paragraph=para)
            related_list_paragraphs = list_paragraphs.filter(document_paragraph=para)
            related_tables = tables.filter(document_paragraph=para)
            related_images = images.filter(document_paragraph=para)
            related_captions = captions.filter(document_paragraph=para)

            # Check elements exist
            has_heading = related_headings.exists()
            has_list_paragraph = related_list_paragraphs.exists()
            has_table = related_tables.exists()
            has_caption = related_captions.exists()
            has_image = related_images.exists()

            # Case heading
            if has_heading:
                for heading in related_headings:
                    add_headings(doc, heading.content, level=1)
                # Subheading check
                related_subheadings = Subheading.objects.filter(heading=heading)
                for subheading in related_subheadings:
                    subheading_level = int(re.search(r'\d+', subheading.type).group()) if re.search(r'\d+', subheading.type) else 2
                    add_headings(doc, subheading.content, level=subheading_level)    

            # Case paragraph stand alone
            add_paragraphs(doc, para.content)

            # Case List paragraph
            if has_list_paragraph:
                for list_para in related_list_paragraphs:
                    add_list_paragraph(doc, list_para.content)

            # Case table
            if has_table:
                for table in related_tables:
                    add_table_to_doc(doc, table)

            # Case Image
            if has_image:
                for image in related_images:
                    add_image_to_doc(doc, image.content)

            # Case Caption
            if has_caption:
                for caption in related_captions:
                    add_caption_to_doc(doc, caption.content)

        # Case Headings without paragraphs ID
        standalone_headings = headings.filter(document_paragraph__isnull=True)
        for heading in standalone_headings:
            add_headings(doc, heading.content, level=1)
        # Subheading
        related_subheadings = Subheading.objects.filter(heading=heading)
        for subheading in related_subheadings:
            subheading_level = int(re.search(r'\d+', subheading.type).group()) if re.search(r'\d+', subheading.type) else 2
            add_headings(doc, subheading.content, level=subheading_level)
            
        # Add endnotes
        for endnote in EndNote.objects.filter(user_document=user_doc):
            add_paragraphs(doc, endnote.content)
        # Save the document
        result = upload_to_s3.upload_to_s3(doc, file_path)
        return f"https://group1-bucket.s3.ap-southeast-1.amazonaws.com/{result}"
